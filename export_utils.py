"""
Enhanced Export Utilities for BRD Generation with Image Embedding Support
Enterprise-grade export functionality with proper image handling
"""

import pandas as pd
import re
import base64
from io import BytesIO
from datetime import datetime
from typing import Dict, Any, List, Tuple
from PIL import Image
from utils.logger import get_logger

logger = get_logger(__name__)

# Optional imports with enhanced fallbacks
try:
    import docx
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. Word export will be disabled.")

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.lib.utils import ImageReader
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logger.warning("ReportLab not available. PDF export will be disabled.")

class ImageProcessor:
    """Process and prepare images for export"""
    
    @staticmethod
    def extract_image_references(content: str) -> List[Tuple[str, str, str, str]]:
        """Extract image references from content"""
        # Pattern: [IMAGE_EMBED:image_id|image_type|description]
        pattern = r'\[IMAGE_EMBED:([^|]+)\|([^|]+)\|([^\]]+)\]'
        matches = re.findall(pattern, content)
        
        references = []
        for match in matches:
            image_id, image_type, description = match
            references.append((image_id.strip(), image_type.strip(), description.strip(), f"[IMAGE_EMBED:{image_id}|{image_type}|{description}]"))
        
        return references
    
    @staticmethod
    def prepare_image_for_export(image_data: str, max_width: int = 600, max_height: int = 400) -> BytesIO:
        """Prepare image for export by resizing and optimizing"""
        try:
            # Handle both string and dict formats
            if isinstance(image_data, dict):
                img_data_b64 = image_data.get('data', image_data)
            else:
                img_data_b64 = image_data
            
            # Decode base64 image
            img_data = base64.b64decode(img_data_b64)
            img = Image.open(BytesIO(img_data))
            
            # Convert to RGB if necessary
            if img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Resize while maintaining aspect ratio
            img.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)
            
            # Save optimized image
            buffer = BytesIO()
            img.save(buffer, format='PNG', optimize=True, quality=90)
            buffer.seek(0)
            
            return buffer
            
        except Exception as e:
            logger.error(f"Error preparing image for export: {e}")
            return None
    
    @staticmethod
    def get_image_caption(image_id: str, image_type: str, description: str) -> str:
        """Generate appropriate caption for exported image"""
        type_labels = {
            'basel_mar_reference': 'Basel MAR Reference',
            'risk_weight_formula': 'Risk Weight Formula',
            'capital_requirement': 'Capital Requirement',
            'correlation_formula': 'Correlation Formula',
            'sensitivity_measure': 'Sensitivity Measure',
            'regulatory_table': 'Regulatory Table',
            'structured_table': 'Data Table',
            'mathematical_expression': 'Mathematical Expression',
            'mathematical_formula': 'Mathematical Formula'
        }
        
        label = type_labels.get(image_type, 'Visual Content')
        return f"Figure: {label} - {description}"

def export_to_word_docx_enhanced(brd_content: Dict[str, Any], images: Dict[str, str] = None) -> BytesIO:
    """Enhanced Word export with proper image embedding"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not available. Please install it to use Word export.")
    
    try:
        doc = docx.Document()
        
        # Enhanced document styling
        styles = doc.styles
        
        # Create custom styles
        try:
            title_style = styles.add_style('CustomTitle', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Arial'
            title_style.font.size = Pt(24)
            title_style.font.bold = True
            title_style.font.color.rgb = docx.shared.RGBColor(0, 51, 102)
        except:
            title_style = styles['Title']
        
        try:
            heading_style = styles.add_style('CustomHeading', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
            heading_style.font.name = 'Arial'
            heading_style.font.size = Pt(16)
            heading_style.font.bold = True
            heading_style.font.color.rgb = docx.shared.RGBColor(0, 102, 204)
        except:
            heading_style = styles['Heading 1']
        
        # Add enhanced title
        title = doc.add_paragraph()
        title.style = title_style
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('Business Requirements Document')
        
        # Add generation info
        info_para = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # Process each section
        for section_name, content in brd_content.items():
            # Add section heading
            heading = doc.add_paragraph()
            heading.style = heading_style
            heading.add_run(section_name)
            
            if isinstance(content, dict):
                # Handle parent sections with subsections
                for subsection_name, subcontent in content.items():
                    # Add subsection heading
                    subheading = doc.add_heading(subsection_name, level=2)
                    
                    if isinstance(subcontent, pd.DataFrame) and not subcontent.empty:
                        # Add enhanced table
                        add_enhanced_table_to_doc(doc, subcontent, subsection_name)
                    else:
                        # Process text content with image embedding
                        add_enhanced_text_to_doc(doc, str(subcontent), images)
                    
                    doc.add_paragraph()  # Add spacing
            
            elif isinstance(content, pd.DataFrame) and not content.empty:
                # Add enhanced table
                add_enhanced_table_to_doc(doc, content, section_name)
            
            else:
                # Process text content with image embedding
                add_enhanced_text_to_doc(doc, str(content), images)
            
            doc.add_page_break()
        
        # Save to bytes
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes
    
    except Exception as e:
        logger.error(f"Error exporting to Word: {e}")
        raise

def add_enhanced_table_to_doc(doc, dataframe: pd.DataFrame, table_name: str):
    """Add enhanced table to Word document"""
    try:
        # Add table caption
        caption = doc.add_paragraph(f"Table: {table_name}")
        caption.style = 'Caption'
        
        # Create table
        table = doc.add_table(rows=1, cols=len(dataframe.columns))
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Add header row with styling
        hdr_cells = table.rows[0].cells
        for i, column in enumerate(dataframe.columns):
            hdr_cells[i].text = str(column)
            # Make header bold
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Add data rows
        for _, row in dataframe.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                cell_text = str(value) if pd.notna(value) else ""
                # Handle long text
                if len(cell_text) > 100:
                    cell_text = cell_text[:97] + "..."
                row_cells[i].text = cell_text
                
                # Adjust cell formatting
                for paragraph in row_cells[i].paragraphs:
                    paragraph.style = 'Normal'
        
        # Set column widths
        for column in table.columns:
            for cell in column.cells:
                cell.width = Inches(1.5)
                
    except Exception as e:
        logger.warning(f"Error adding table {table_name}: {e}")
        # Fallback to simple table
        doc.add_paragraph(f"Table: {table_name} (Simplified)")
        doc.add_paragraph(str(dataframe.to_string()))

def add_enhanced_text_to_doc(doc, content: str, images: Dict[str, str] = None):
    """Add text content with embedded images to Word document"""
    if not content:
        return
    
    # Extract image references
    image_refs = ImageProcessor.extract_image_references(content)
    
    if not image_refs or not images:
        # No images to embed, add as regular text
        clean_content = re.sub(r'\[IMAGE_EMBED:[^\]]+\]', '[Visual Content Reference]', content)
        paragraphs = clean_content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
        return
    
    # Split content by image references and process
    current_content = content
    
    for image_id, image_type, description, full_ref in image_refs:
        if image_id in images:
            # Split content at this image reference
            parts = current_content.split(full_ref, 1)
            
            # Add text before image
            if parts[0].strip():
                text_paragraphs = parts[0].strip().split('\n\n')
                for para_text in text_paragraphs:
                    if para_text.strip():
                        doc.add_paragraph(para_text.strip())
            
            # Add image
            try:
                image_buffer = ImageProcessor.prepare_image_for_export(images[image_id])
                if image_buffer:
                    # Add image caption
                    caption = ImageProcessor.get_image_caption(image_id, image_type, description)
                    doc.add_paragraph(caption, style='Caption')
                    
                    # Add image
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    run.add_picture(image_buffer, width=Inches(5))
                    
                    # Add spacing
                    doc.add_paragraph()
                    
                else:
                    doc.add_paragraph(f"[Image: {description}]", style='Caption')
                    
            except Exception as e:
                logger.warning(f"Error embedding image {image_id}: {e}")
                doc.add_paragraph(f"[Image: {description}]", style='Caption')
            
            # Continue with remaining content
            current_content = parts[1] if len(parts) > 1 else ""
        else:
            # Image not found, replace with placeholder
            current_content = current_content.replace(full_ref, f"[Image Reference: {description}]")
    
    # Add any remaining content
    if current_content.strip():
        remaining_paragraphs = current_content.strip().split('\n\n')
        for para_text in remaining_paragraphs:
            if para_text.strip():
                doc.add_paragraph(para_text.strip())

def export_to_pdf_enhanced(brd_content: Dict[str, Any], images: Dict[str, str] = None) -> BytesIO:
    """Enhanced PDF export with proper image embedding"""
    if not REPORTLAB_AVAILABLE:
        raise ImportError("ReportLab not available. Please install it to use PDF export.")
    
    try:
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=1*inch, bottomMargin=1*inch)
        story = []
        styles = getSampleStyleSheet()
        
        # Enhanced custom styles
        title_style = ParagraphStyle(
            'EnhancedTitle',
            parent=styles['Title'],
            fontSize=26,
            textColor=colors.Color(0, 0.2, 0.4),
            alignment=1,  # Center alignment
            spaceAfter=30,
            fontName='Helvetica-Bold'
        )
        
        heading_style = ParagraphStyle(
            'EnhancedHeading',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.Color(0, 0.4, 0.8),
            spaceAfter=15,
            spaceBefore=20,
            fontName='Helvetica-Bold'
        )
        
        subheading_style = ParagraphStyle(
            'EnhancedSubHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.Color(0.2, 0.2, 0.2),
            spaceAfter=12,
            spaceBefore=15,
            fontName='Helvetica-Bold'
        )
        
        caption_style = ParagraphStyle(
            'ImageCaption',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.Color(0.3, 0.3, 0.3),
            alignment=1,  # Center alignment
            spaceAfter=10,
            spaceBefore=5,
            fontName='Helvetica-Oblique'
        )
        
        # Add enhanced title
        story.append(Paragraph("Business Requirements Document", title_style))
        story.append(Paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", styles['Normal']))
        story.append(Spacer(1, 30))
        
        # Process each section
        for section_name, content in brd_content.items():
            # Add section heading
            story.append(Paragraph(section_name, heading_style))
            
            if isinstance(content, dict):
                # Handle subsections
                for subsection_name, subcontent in content.items():
                    story.append(Paragraph(subsection_name, subheading_style))
                    
                    if isinstance(subcontent, pd.DataFrame) and not subcontent.empty:
                        # Add enhanced table
                        add_enhanced_table_to_pdf(story, subcontent, subsection_name, styles)
                    else:
                        # Add text with images
                        add_enhanced_text_to_pdf(story, str(subcontent), images, styles, caption_style)
                    
                    story.append(Spacer(1, 15))
            
            elif isinstance(content, pd.DataFrame) and not content.empty:
                # Add enhanced table
                add_enhanced_table_to_pdf(story, content, section_name, styles)
            
            else:
                # Add text with images
                add_enhanced_text_to_pdf(story, str(content), images, styles, caption_style)
            
            story.append(Spacer(1, 25))
        
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    except Exception as e:
        logger.error(f"Error exporting to PDF: {e}")
        raise

def add_enhanced_table_to_pdf(story: List, dataframe: pd.DataFrame, table_name: str, styles):
    """Add enhanced table to PDF"""
    try:
        # Add table caption
        story.append(Paragraph(f"Table: {table_name}", styles['Caption']))
        
        # Prepare table data with size limits
        table_data = []
        
        # Add headers
        headers = [str(col)[:25] + "..." if len(str(col)) > 25 else str(col) for col in dataframe.columns]
        table_data.append(headers)
        
        # Add data rows (limit to prevent page overflow)
        max_rows = min(20, len(dataframe))
        for _, row in dataframe.head(max_rows).iterrows():
            row_data = []
            for val in row:
                cell_text = str(val) if pd.notna(val) else ""
                # Truncate long text
                if len(cell_text) > 30:
                    cell_text = cell_text[:27] + "..."
                row_data.append(cell_text)
            table_data.append(row_data)
        
        # Create table with enhanced styling
        table = Table(table_data, repeatRows=1)
        table.setStyle(TableStyle([
            # Header styling
            ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.8, 0.8, 0.9)),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.Color(0.1, 0.1, 0.4)),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            
            # Data styling
            ('BACKGROUND', (0, 1), (-1, -1), colors.Color(0.95, 0.95, 0.98)),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 1, colors.Color(0.7, 0.7, 0.8)),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            
            # Alternating row colors
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.Color(0.95, 0.95, 0.98), colors.Color(0.98, 0.98, 1.0)])
        ]))
        
        story.append(table)
        
        # Add note if table was truncated
        if len(dataframe) > max_rows:
            story.append(Paragraph(f"Note: Table truncated to {max_rows} rows of {len(dataframe)} total.", styles['Caption']))
        
    except Exception as e:
        logger.warning(f"Error adding table {table_name} to PDF: {e}")
        # Fallback
        story.append(Paragraph(f"Table: {table_name} (Error in formatting)", styles['Caption']))

def add_enhanced_text_to_pdf(story: List, content: str, images: Dict[str, str], styles, caption_style):
    """Add text content with embedded images to PDF"""
    if not content:
        return
    
    # Extract image references
    image_refs = ImageProcessor.extract_image_references(content)
    
    if not image_refs or not images:
        # No images to embed, add as regular text
        clean_content = re.sub(r'\[IMAGE_EMBED:[^\]]+\]', '[Visual Content Reference]', content)
        paragraphs = clean_content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                story.append(Paragraph(para_text.strip(), styles['Normal']))
                story.append(Spacer(1, 6))
        return
    
    # Process content with images
    current_content = content
    
    for image_id, image_type, description, full_ref in image_refs:
        if image_id in images:
            # Split content at this image reference
            parts = current_content.split(full_ref, 1)
            
            # Add text before image
            if parts[0].strip():
                text_paragraphs = parts[0].strip().split('\n\n')
                for para_text in text_paragraphs:
                    if para_text.strip():
                        story.append(Paragraph(para_text.strip(), styles['Normal']))
                        story.append(Spacer(1, 6))
            
            # Add image
            try:
                image_buffer = ImageProcessor.prepare_image_for_export(images[image_id], max_width=500, max_height=300)
                if image_buffer:
                    # Add image
                    rl_image = RLImage(image_buffer, width=4*inch, height=3*inch, kind='proportional')
                    story.append(rl_image)
                    
                    # Add caption
                    caption = ImageProcessor.get_image_caption(image_id, image_type, description)
                    story.append(Paragraph(caption, caption_style))
                    story.append(Spacer(1, 15))
                else:
                    story.append(Paragraph(f"[Image: {description}]", caption_style))
                    story.append(Spacer(1, 10))
                    
            except Exception as e:
                logger.warning(f"Error embedding image {image_id} in PDF: {e}")
                story.append(Paragraph(f"[Image: {description}]", caption_style))
                story.append(Spacer(1, 10))
            
            # Continue with remaining content
            current_content = parts[1] if len(parts) > 1 else ""
        else:
            # Image not found, replace with placeholder
            current_content = current_content.replace(full_ref, f"[Image Reference: {description}]")
    
    # Add any remaining content
    if current_content.strip():
        remaining_paragraphs = current_content.strip().split('\n\n')
        for para_text in remaining_paragraphs:
            if para_text.strip():
                story.append(Paragraph(para_text.strip(), styles['Normal']))
                story.append(Spacer(1, 6))

def export_to_excel_enhanced(brd_content: Dict[str, Any], images: Dict[str, str] = None) -> BytesIO:
    """Enhanced Excel export with image references and metadata"""
    try:
        buffer = BytesIO()
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            sheet_count = 1
            image_summary = []
            
            for section_name, content in brd_content.items():
                if isinstance(content, dict):
                    for subsection_name, subcontent in content.items():
                        if isinstance(subcontent, pd.DataFrame) and not subcontent.empty:
                            # Create safe sheet name
                            sheet_name = f"Sheet{sheet_count}_{subsection_name.split('.')[0]}"[:31]
                            sheet_name = re.sub(r'[^\w\s-]', '', sheet_name).strip()[:31]
                            
                            try:
                                subcontent.to_excel(writer, sheet_name=sheet_name, index=False)
                                sheet_count += 1
                            except Exception as e:
                                logger.warning(f"Error writing sheet {sheet_name}: {e}")
                        else:
                            # Extract image references for summary
                            content_str = str(subcontent)
                            image_refs = ImageProcessor.extract_image_references(content_str)
                            for image_id, image_type, description, _ in image_refs:
                                image_summary.append({
                                    'Section': subsection_name,
                                    'Image_ID': image_id,
                                    'Type': image_type,
                                    'Description': description
                                })
                                
                elif isinstance(content, pd.DataFrame) and not content.empty:
                    # Create safe sheet name  
                    sheet_name = f"Sheet{sheet_count}_{section_name.split('.')[0]}"[:31]
                    sheet_name = re.sub(r'[^\w\s-]', '', sheet_name).strip()[:31]
                    
                    try:
                        content.to_excel(writer, sheet_name=sheet_name, index=False)
                        sheet_count += 1
                    except Exception as e:
                        logger.warning(f"Error writing sheet {sheet_name}: {e}")
                else:
                    # Extract image references for summary
                    content_str = str(content)
                    image_refs = ImageProcessor.extract_image_references(content_str)
                    for image_id, image_type, description, _ in image_refs:
                        image_summary.append({
                            'Section': section_name,
                            'Image_ID': image_id,
                            'Type': image_type,
                            'Description': description
                        })
            
            # Add image summary sheet
            if image_summary:
                image_df = pd.DataFrame(image_summary)
                try:
                    image_df.to_excel(writer, sheet_name='Image_References', index=False)
                except Exception as e:
                    logger.warning(f"Error writing image summary sheet: {e}")
        
        buffer.seek(0)
        return buffer
    
    except Exception as e:
        logger.error(f"Error exporting to Excel: {e}")
        raise

def export_to_json_enhanced(brd_content: Dict[str, Any], images: Dict[str, str] = None) -> str:
    """Enhanced JSON export with image metadata"""
    try:
        import json
        
        # Convert DataFrames and add image metadata
        json_content = {}
        image_metadata = {}
        
        for section_name, content in brd_content.items():
            if isinstance(content, dict):
                json_content[section_name] = {}
                for subsection_name, subcontent in content.items():
                    if isinstance(subcontent, pd.DataFrame):
                        json_content[section_name][subsection_name] = subcontent.to_dict('records')
                    else:
                        content_str = str(subcontent)
                        json_content[section_name][subsection_name] = content_str
                        
                        # Extract image references
                        image_refs = ImageProcessor.extract_image_references(content_str)
                        if image_refs:
                            image_metadata[subsection_name] = [
                                {
                                    'image_id': img_id,
                                    'type': img_type,
                                    'description': desc
                                }
                                for img_id, img_type, desc, _ in image_refs
                            ]
            elif isinstance(content, pd.DataFrame):
                json_content[section_name] = content.to_dict('records')
            else:
                content_str = str(content)
                json_content[section_name] = content_str
                
                # Extract image references
                image_refs = ImageProcessor.extract_image_references(content_str)
                if image_refs:
                    image_metadata[section_name] = [
                        {
                            'image_id': img_id,
                            'type': img_type,
                            'description': desc
                        }
                        for img_id, img_type, desc, _ in image_refs
                    ]
        
        # Add metadata
        export_data = {
            'brd_content': json_content,
            'image_metadata': image_metadata,
            'export_info': {
                'generated_at': datetime.now().isoformat(),
                'total_sections': len(json_content),
                'total_image_references': sum(len(refs) for refs in image_metadata.values()),
                'available_images': len(images) if images else 0
            }
        }
        
        return json.dumps(export_data, indent=2, ensure_ascii=False)
    
    except Exception as e:
        logger.error(f"Error exporting to JSON: {e}")
        raise

# Enhanced export function mappings
def export_to_word_docx(brd_content: Dict[str, Any], images: Dict[str, str] = None) -> BytesIO:
    """Enhanced Word export - main entry point"""
    return export_to_word_docx_enhanced(brd_content, images)

def export_to_pdf(brd_content: Dict[str, Any], images: Dict[str, str] = None) -> BytesIO:
    """Enhanced PDF export - main entry point"""
    return export_to_pdf_enhanced(brd_content, images)

def export_to_excel(brd_content: Dict[str, Any], images: Dict[str, str] = None) -> BytesIO:
    """Enhanced Excel export - main entry point"""
    return export_to_excel_enhanced(brd_content, images)

def export_to_json(brd_content: Dict[str, Any], images: Dict[str, str] = None) -> str:
    """Enhanced JSON export - main entry point"""
    return export_to_json_enhanced(brd_content, images)

# Additional utility functions for enhanced export functionality

def validate_export_content(brd_content: Dict[str, Any], images: Dict[str, str] = None) -> Dict[str, Any]:
    """Validate content before export and provide recommendations"""
    validation_results = {
        'content_valid': True,
        'warnings': [],
        'errors': [],
        'statistics': {},
        'recommendations': []
    }
    
    try:
        # Content validation
        total_sections = len(brd_content)
        empty_sections = 0
        sections_with_images = 0
        total_image_refs = 0
        
        for section_name, content in brd_content.items():
            # Check for empty content
            if not content or (isinstance(content, str) and not content.strip()):
                empty_sections += 1
                validation_results['warnings'].append(f"Empty section: {section_name}")
            
            # Check for image references
            content_str = str(content)
            image_refs = ImageProcessor.extract_image_references(content_str)
            if image_refs:
                sections_with_images += 1
                total_image_refs += len(image_refs)
                
                # Validate image availability
                for image_id, _, _, _ in image_refs:
                    if not images or image_id not in images:
                        validation_results['errors'].append(f"Missing image: {image_id} in {section_name}")
                        validation_results['content_valid'] = False
        
        # Statistics
        validation_results['statistics'] = {
            'total_sections': total_sections,
            'empty_sections': empty_sections,
            'sections_with_images': sections_with_images,
            'total_image_references': total_image_refs,
            'available_images': len(images) if images else 0,
            'completion_rate': (total_sections - empty_sections) / total_sections * 100 if total_sections > 0 else 0
        }
        
        # Recommendations
        if empty_sections > 0:
            validation_results['recommendations'].append(f"Complete {empty_sections} empty sections before export")
        
        if total_image_refs > 0 and not images:
            validation_results['recommendations'].append("Images referenced but not available - consider re-processing document")
        
        if validation_results['statistics']['completion_rate'] < 80:
            validation_results['recommendations'].append("Consider completing more sections for comprehensive BRD")
        
        if sections_with_images > 0:
            validation_results['recommendations'].append(f"Enhanced export recommended - {sections_with_images} sections contain visual content")
    
    except Exception as e:
        validation_results['errors'].append(f"Validation error: {str(e)}")
        validation_results['content_valid'] = False
    
    return validation_results

def generate_export_summary(brd_content: Dict[str, Any], images: Dict[str, str] = None) -> str:
    """Generate summary of export content"""
    try:
        validation = validate_export_content(brd_content, images)
        stats = validation['statistics']
        
        summary = f"""
BRD EXPORT SUMMARY
==================

Content Overview:
- Total Sections: {stats['total_sections']}
- Completion Rate: {stats['completion_rate']:.1f}%
- Sections with Visual Content: {stats['sections_with_images']}
- Total Image References: {stats['total_image_references']}
- Available Images: {stats['available_images']}

Export Readiness:
- Content Valid: {'Yes' if validation['content_valid'] else 'No'}
- Warnings: {len(validation['warnings'])}
- Errors: {len(validation['errors'])}

"""
        
        if validation['warnings']:
            summary += "\nWarnings:\n"
            for warning in validation['warnings']:
                summary += f"- {warning}\n"
        
        if validation['errors']:
            summary += "\nErrors:\n"
            for error in validation['errors']:
                summary += f"- {error}\n"
        
        if validation['recommendations']:
            summary += "\nRecommendations:\n"
            for rec in validation['recommendations']:
                summary += f"- {rec}\n"
        
        summary += f"\nGenerated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        
        return summary
        
    except Exception as e:
        return f"Error generating export summary: {str(e)}"

def create_export_package(brd_content: Dict[str, Any], images: Dict[str, str] = None, formats: List[str] = None) -> Dict[str, BytesIO]:
    """Create a package of exports in multiple formats"""
    if formats is None:
        formats = ['word', 'pdf', 'excel', 'json']
    
    export_package = {}
    
    try:
        if 'word' in formats:
            try:
                export_package['word'] = export_to_word_docx_enhanced(brd_content, images)
                logger.info("Word export completed successfully")
            except Exception as e:
                logger.error(f"Word export failed: {e}")
                export_package['word_error'] = str(e)
        
        if 'pdf' in formats:
            try:
                export_package['pdf'] = export_to_pdf_enhanced(brd_content, images)
                logger.info("PDF export completed successfully")
            except Exception as e:
                logger.error(f"PDF export failed: {e}")
                export_package['pdf_error'] = str(e)
        
        if 'excel' in formats:
            try:
                export_package['excel'] = export_to_excel_enhanced(brd_content, images)
                logger.info("Excel export completed successfully")
            except Exception as e:
                logger.error(f"Excel export failed: {e}")
                export_package['excel_error'] = str(e)
        
        if 'json' in formats:
            try:
                json_content = export_to_json_enhanced(brd_content, images)
                json_buffer = BytesIO(json_content.encode('utf-8'))
                export_package['json'] = json_buffer
                logger.info("JSON export completed successfully")
            except Exception as e:
                logger.error(f"JSON export failed: {e}")
                export_package['json_error'] = str(e)
        
        # Add export summary
        try:
            summary = generate_export_summary(brd_content, images)
            summary_buffer = BytesIO(summary.encode('utf-8'))
            export_package['summary'] = summary_buffer
        except Exception as e:
            logger.error(f"Summary generation failed: {e}")
    
    except Exception as e:
        logger.error(f"Export package creation failed: {e}")
        export_package['package_error'] = str(e)
    
    return export_package
